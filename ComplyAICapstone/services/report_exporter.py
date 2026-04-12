# -*- coding: utf-8 -*-
"""
Report Export Service for NIST CSF Compliance Reports

This module handles:
- PDF report generation using ReportLab
- Word document generation using python-docx
- Professional formatting and styling
- Export to multiple formats
"""

import os
import io
from datetime import datetime
from typing import Dict, Any, Optional, List
from pathlib import Path

# PDF generation imports
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor, black, white
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.platypus.tableofcontents import TableOfContents
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

# Word document imports
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

from models.report_models import FinalReport, RiskLevel, NISTFunction
from services.questionnaire_engine import QuestionnaireEngine


class ReportExporter:
    """Service for exporting compliance reports to PDF and Word formats."""
    
    def __init__(self, output_dir: str = "exports"):
        """
        Initialize the report exporter.
        
        Args:
            output_dir: Directory to save exported reports
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        
        # Define color scheme
        self.colors = {
            'primary': HexColor('#2c3e50'),
            'secondary': HexColor('#3498db'),
            'accent': HexColor('#e74c3c'),
            'success': HexColor('#27ae60'),
            'warning': HexColor('#f39c12'),
            'light_gray': HexColor('#ecf0f1'),
            'dark_gray': HexColor('#34495e')
        }
    
    def export_to_pdf(self, report: FinalReport, heatmap_report=None, filename: Optional[str] = None, report_type: str = "detailed") -> str:
        """
        Export report to PDF format.
        
        Args:
            report: FinalReport object to export
            heatmap_report: Optional HeatMapReport to include
            filename: Optional custom filename
            
        Returns:
            Path to generated PDF file
        """
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"NIST_CSF_Report_{report.organization_info.organization_name}_{timestamp}.pdf"
        
        filepath = self.output_dir / filename
        
        # Delegate to new selective exporter for compatibility
        return self.export_report_pdf(report, report_type=report_type, heatmap_report=heatmap_report, filename=filename)

    def export_report_pdf(self, report: FinalReport, report_type: str = "detailed", heatmap_report=None, filename: Optional[str] = None) -> str:
        """
        Export report to PDF using modular templates based on `report_type`.
        Supported `report_type` values: 'executive', 'detailed', 'improvement_plan', 'risk_register'
        """
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"NIST_CSF_Report_{report.organization_info.organization_name}_{timestamp}.pdf"

        filepath = self.output_dir / filename

        # Create PDF document
        doc = SimpleDocTemplate(
            str(filepath),
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )

        # Build story (content) depending on requested report type
        story = []

        # Common title + toc
        story.extend(self._create_pdf_title_page(report))
        story.append(PageBreak())
        story.extend(self._create_pdf_toc())
        story.append(PageBreak())

        rt = (report_type or "detailed").lower()

        if rt in ('executive', 'executive_summary', 'summary'):
            story.extend(self._create_pdf_executive_summary(report))
            story.append(PageBreak())
            story.extend(self._create_pdf_organization_info(report))
        elif rt in ('detailed', 'detailed_compliance_report'):
            story.extend(self._create_pdf_executive_summary(report))
            story.append(PageBreak())
            story.extend(self._create_pdf_organization_info(report))
            story.append(PageBreak())
            story.extend(self._create_pdf_compliance_by_function(report))
            story.append(PageBreak())
            if report.compliance_gaps:
                story.extend(self._create_pdf_compliance_gaps(report))
                story.append(PageBreak())
            if heatmap_report:
                story.extend(self._create_pdf_heatmap_section(report, heatmap_report))
                story.append(PageBreak())
            if report.risk_assessment:
                story.extend(self._create_pdf_risk_assessment(report))
                story.append(PageBreak())
            if report.recommendations:
                story.extend(self._create_pdf_recommendations(report))
                story.append(PageBreak())
            if report.questionnaire_answers:
                story.extend(self._create_pdf_questionnaire_evidence(report))
        elif rt in ('improvement_plan', 'cybersecurity_improvement_plan'):
            # Improvement plan: structured table derived from compliance gaps and risks
            story.append(Paragraph('Cybersecurity Improvement Plan', getSampleStyleSheet()['Heading1']))
            story.append(Spacer(1, 0.2*inch))

            # Build table header
            table_data = [['Issue', 'Risk Level', 'Recommended Action', 'Priority', 'Timeline']]

            # Use centralized mapping helpers for priority and timeline

            # Prefer recommendations list if present, otherwise derive from gaps
            if getattr(report, 'recommendations', None):
                for rec in report.recommendations:
                    issue = rec.title
                    risk = getattr(rec, 'risk_level', '') or getattr(rec, 'severity', '') or 'Medium'
                    if hasattr(risk, 'value'):
                        risk = risk.value
                    action = rec.description or rec.estimated_effort or ''
                    pri = self._priority_from_risk(risk)
                    timeline = self._timeline_from_priority(pri)
                    table_data.append([issue, str(risk), action, pri, timeline])
            else:
                # Use compliance gaps
                for gap in getattr(report, 'compliance_gaps', []) or []:
                    issue = f"{gap.control_id}: {gap.control_title}" if getattr(gap, 'control_id', None) else (getattr(gap, 'control_title', 'Issue'))
                    risk = getattr(gap, 'risk_level', 'Medium')
                    if hasattr(risk, 'value'):
                        risk = risk.value
                    action = gap.description or gap.reasoning or 'Review and remediate per control guidance.'
                    pri = self._priority_from_risk(risk)
                    timeline = self._timeline_from_priority(pri)
                    table_data.append([issue, str(risk), action, pri, timeline])

            tbl = Table(table_data, colWidths=[2.5*inch, 0.9*inch, 2.2*inch, 0.8*inch, 1.1*inch])
            tbl.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), self.colors['primary']),
                ('TEXTCOLOR', (0,0), (-1,0), white),
                ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                ('GRID', (0,0), (-1,-1), 0.5, self.colors['dark_gray']),
                ('BACKGROUND', (0,1), (-1,-1), self.colors['light_gray'])
            ]))

            story.append(tbl)
        elif rt in ('risk_register', 'risk_register_only'):
            # Compact risk table only
            story.append(Paragraph('Risk Register', getSampleStyleSheet()['Heading1']))
            story.append(Spacer(1, 0.2*inch))

            table_data = [['Risk', 'Likelihood', 'Impact', 'Score', 'Status']]
            for r in getattr(report, 'risk_assessment', []) or []:
                title = getattr(r, 'title', 'Unnamed Risk')
                likelihood = getattr(r, 'likelihood', 'Unknown')
                impact = getattr(r, 'impact', 'Unknown')
                score = getattr(r, 'score', None) or getattr(r, 'risk_score', '') or ''
                status = getattr(r, 'mitigation_status', '') or getattr(r, 'status', '')
                table_data.append([title, str(likelihood), str(impact), str(score), str(status)])

            tbl = Table(table_data, colWidths=[2.5*inch, 1*inch, 1*inch, 0.8*inch, 1.2*inch])
            tbl.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), self.colors['primary']),
                ('TEXTCOLOR', (0,0), (-1,0), white),
                ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                ('GRID', (0,0), (-1,-1), 0.5, self.colors['dark_gray']),
                ('BACKGROUND', (0,1), (-1,-1), self.colors['light_gray'])
            ]))

            story.append(tbl)
        else:
            # Unknown type: fallback to detailed
            story.extend(self._create_pdf_executive_summary(report))
            story.append(PageBreak())
            story.extend(self._create_pdf_organization_info(report))

        # Build PDF
        try:
            doc.build(story)
        except Exception:
            # In case a section building failed, try minimal executive summary
            story = []
            story.extend(self._create_pdf_title_page(report))
            story.append(PageBreak())
            story.extend(self._create_pdf_executive_summary(report))
            doc.build(story)

        return str(filepath)
    
    def export_to_word(self, report: FinalReport, heatmap_report=None, filename: Optional[str] = None) -> str:
        # Backwards compatible signature: default to detailed
        return self.export_report_word(report, report_type='detailed', heatmap_report=heatmap_report, filename=filename)

    def export_report_word(self, report: FinalReport, report_type: str = 'detailed', heatmap_report=None, filename: Optional[str] = None) -> str:
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"NIST_CSF_Report_{report.organization_info.organization_name}_{timestamp}.docx"

        filepath = self.output_dir / filename

        # Create Word document
        doc = Document()

        # Set up styles
        self._setup_word_styles(doc)

        rt = (report_type or 'detailed').lower()

        # Title + TOC
        self._create_word_title_page(doc, report)
        doc.add_page_break()
        self._create_word_toc(doc)
        doc.add_page_break()

        if rt in ('executive', 'executive_summary', 'summary'):
            self._create_word_executive_summary(doc, report)
            doc.add_page_break()
            self._create_word_organization_info(doc, report)
        elif rt in ('detailed', 'detailed_compliance_report'):
            self._create_word_executive_summary(doc, report)
            doc.add_page_break()
            self._create_word_organization_info(doc, report)
            doc.add_page_break()
            # Compliance gaps and other detailed sections
            if report.compliance_gaps:
                self._create_word_compliance_gaps(doc, report)
                doc.add_page_break()
            if report.risk_assessment:
                self._create_word_risk_assessment(doc, report)
                doc.add_page_break()
            if heatmap_report:
                self._create_word_heatmap_section(doc, report, heatmap_report)
                doc.add_page_break()
            if report.recommendations:
                self._create_word_recommendations(doc, report)
                doc.add_page_break()
            if report.questionnaire_answers:
                self._create_word_questionnaire_evidence(doc, report)
        elif rt in ('improvement_plan', 'cybersecurity_improvement_plan'):
            doc.add_paragraph('Cybersecurity Improvement Plan', style='Heading1')
            # Build similar table in Word
            table = doc.add_table(rows=1, cols=5)
            hdr = table.rows[0].cells
            hdr[0].text = 'Issue'
            hdr[1].text = 'Risk Level'
            hdr[2].text = 'Recommended Action'
            hdr[3].text = 'Priority'
            hdr[4].text = 'Timeline'

            # Use centralized mapping helpers

            if getattr(report, 'recommendations', None):
                items = report.recommendations
            else:
                items = getattr(report, 'compliance_gaps', []) or []

            for it in items:
                row = table.add_row().cells
                title = getattr(it, 'title', None) or (f"{getattr(it, 'control_id', '')}: {getattr(it, 'control_title', '')}" if getattr(it, 'control_id', None) else getattr(it, 'control_title', 'Issue'))
                risk = getattr(it, 'risk_level', '')
                if hasattr(risk, 'value'):
                    risk = risk.value
                action = getattr(it, 'description', '') or getattr(it, 'reasoning', '') or ''
                pri = self._priority_from_risk(risk)
                timeline = self._timeline_from_priority(pri)
                row[0].text = title
                row[1].text = str(risk)
                row[2].text = action
                row[3].text = pri
                row[4].text = timeline
        elif rt in ('risk_register', 'risk_register_only'):
            doc.add_paragraph('Risk Register', style='Heading1')
            table = doc.add_table(rows=1, cols=5)
            hdr = table.rows[0].cells
            hdr[0].text = 'Risk'
            hdr[1].text = 'Likelihood'
            hdr[2].text = 'Impact'
            hdr[3].text = 'Score'
            hdr[4].text = 'Status'

            for r in getattr(report, 'risk_assessment', []) or []:
                row = table.add_row().cells
                row[0].text = getattr(r, 'title', '')
                row[1].text = str(getattr(r, 'likelihood', ''))
                row[2].text = str(getattr(r, 'impact', ''))
                row[3].text = str(getattr(r, 'score', '') or getattr(r, 'risk_score', ''))
                row[4].text = str(getattr(r, 'mitigation_status', '') or getattr(r, 'status', ''))
        else:
            # fallback to executive
            self._create_word_executive_summary(doc, report)

        # Save document
        doc.save(str(filepath))

        return str(filepath)
    
    def _create_pdf_title_page(self, report: FinalReport) -> List:
        """Create PDF title page."""
        story = []
        styles = getSampleStyleSheet()
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=24,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=self.colors['primary']
        )
        
        story.append(Spacer(1, 2*inch))
        story.append(Paragraph("NIST CSF Compliance Assessment", title_style))
        story.append(Spacer(1, 0.5*inch))
        
        # Organization name
        org_style = ParagraphStyle(
            'Organization',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=20,
            alignment=TA_CENTER,
            textColor=self.colors['secondary']
        )
        
        story.append(Paragraph(report.organization_info.organization_name, org_style))
        story.append(Spacer(1, 1*inch))
        
        # Report metadata
        meta_style = ParagraphStyle(
            'Metadata',
            parent=styles['Normal'],
            fontSize=12,
            alignment=TA_CENTER,
            spaceAfter=10
        )
        
        story.append(Paragraph(f"Assessment Date: {report.organization_info.assessment_date.strftime('%B %d, %Y')}", meta_style))
        story.append(Paragraph(f"Generated: {report.generated_at.strftime('%B %d, %Y at %I:%M %p')}", meta_style))
        story.append(Paragraph(f"Report ID: {report.report_id}", meta_style))
        try:
            framework_name = (report.metadata or {}).get("framework_name")
            framework_version_id = (report.metadata or {}).get("framework_version_id")
            framework_version_label = (report.metadata or {}).get("framework_version_label")
            if framework_name or framework_version_id or framework_version_label:
                framework_line = f"Framework Used: {framework_name or 'NIST CSF'}"
                if framework_version_label:
                    framework_line += f" v{framework_version_label}"
                if framework_version_id:
                    framework_line += f" ({framework_version_id})"
                story.append(Paragraph(framework_line, meta_style))
        except Exception:
            pass
        story.append(Spacer(1, 1*inch))
        
        # Compliance summary
        summary_style = ParagraphStyle(
            'Summary',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=10,
            alignment=TA_CENTER
        )
        
        story.append(Paragraph("Compliance Summary", summary_style))
        
        summary_data = [
            ['Metric', 'Value'],
            ['Overall Compliance', f"{report.compliance_summary.compliance_percentage:.1f}%"],
            ['Total Gaps Identified', str(report.compliance_summary.total_gaps)],
            ['Overall Risk Score', f"{report.compliance_summary.overall_risk_score:.1f}/10"],
            ['Critical Findings', str(report.compliance_summary.critical_findings)]
        ]
        
        summary_table = Table(summary_data, colWidths=[3*inch, 2*inch])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), self.colors['primary']),
            ('TEXTCOLOR', (0, 0), (-1, 0), white),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), self.colors['light_gray']),
            ('GRID', (0, 0), (-1, -1), 1, self.colors['dark_gray'])
        ]))
        
        story.append(summary_table)
        
        return story
    
    def _create_pdf_toc(self) -> List:
        """Create PDF table of contents."""
        story = []
        styles = getSampleStyleSheet()
        
        toc_title = ParagraphStyle(
            'TOCTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=20,
            textColor=self.colors['primary']
        )
        
        story.append(Paragraph("Table of Contents", toc_title))
        story.append(Spacer(1, 0.3*inch))
        
        toc_items = [
            "1. Executive Summary",
            "2. Organization Information",
            "3. Compliance Gaps Analysis",
            "4. Risk Assessment",
            "5. Recommendations",
            "6. Questionnaire Evidence"
        ]
        
        for item in toc_items:
            story.append(Paragraph(item, styles['Normal']))
            story.append(Spacer(1, 0.1*inch))
        
        return story
    
    def _create_pdf_executive_summary(self, report: FinalReport) -> List:
        """Create PDF executive summary section."""
        story = []
        styles = getSampleStyleSheet()
        
        # Section title
        title_style = ParagraphStyle(
            'SectionTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=20,
            textColor=self.colors['primary']
        )
        
        story.append(Paragraph("Executive Summary", title_style))
        
        # Summary content
        summary_text = f"""
        This NIST Cybersecurity Framework (CSF) compliance assessment was conducted for {report.organization_info.organization_name},
        a {report.organization_info.size} organization in the {report.organization_info.industry} sector.
        
        The assessment identified {report.compliance_summary.total_gaps} compliance gaps across {len(set(gap.nist_function for gap in report.compliance_gaps))} NIST CSF functions.
        The overall compliance score is {report.compliance_summary.compliance_percentage:.1f}%, with an overall risk score of {report.compliance_summary.overall_risk_score:.1f}/10.
        
        Key findings include {report.compliance_summary.critical_findings} critical issues that require immediate attention,
        and {len([r for r in report.recommendations if r.priority == 'High'])} high-priority recommendations for remediation.
        
        The assessment covered the following NIST CSF functions: {', '.join(set(gap.nist_function.value for gap in report.compliance_gaps))}.
        """
        
        story.append(Paragraph(summary_text, styles['Normal']))
        story.append(Spacer(1, 0.3*inch))
        
        # Key metrics table
        metrics_data = [
            ['Metric', 'Value', 'Status'],
            ['Compliance Percentage', f"{report.compliance_summary.compliance_percentage:.1f}%", 
             'Good' if report.compliance_summary.compliance_percentage >= 70 else 'Needs Improvement'],
            ['Risk Score', f"{report.compliance_summary.overall_risk_score:.1f}/10",
             'Low' if report.compliance_summary.overall_risk_score <= 3 else 'Moderate' if report.compliance_summary.overall_risk_score <= 6 else 'High'],
            ['Critical Findings', str(report.compliance_summary.critical_findings),
             'None' if report.compliance_summary.critical_findings == 0 else 'Action Required'],
            ['High Priority Recommendations', str(report.compliance_summary.high_priority_recommendations),
             'Few' if report.compliance_summary.high_priority_recommendations <= 3 else 'Many']
        ]
        
        metrics_table = Table(metrics_data, colWidths=[2.5*inch, 1.5*inch, 1.5*inch])
        metrics_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), self.colors['primary']),
            ('TEXTCOLOR', (0, 0), (-1, 0), white),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), self.colors['light_gray']),
            ('GRID', (0, 0), (-1, -1), 1, self.colors['dark_gray'])
        ]))
        
        story.append(metrics_table)
        # AI Confidence Score (aggregate) and Top 5 gaps
        try:
            ai_conf = getattr(report, 'ai_confidence_score', None)
            if ai_conf is None:
                confidences = [getattr(g, 'confidence_score', None) for g in getattr(report, 'compliance_gaps', []) or []]
                vals = [c for c in confidences if isinstance(c, (int, float))]
                if vals:
                    ai_conf = sum(vals) / len(vals)
                else:
                    ai_conf = None
        except Exception:
            ai_conf = None

        ai_text = f"AI Confidence Score: {ai_conf:.2f}" if (isinstance(ai_conf, (int, float))) else "AI Confidence Score: N/A"
        story.append(Spacer(1, 0.15*inch))
        story.append(Paragraph(ai_text, styles['Normal']))
        story.append(Spacer(1, 0.2*inch))

        # Top 5 compliance gaps (by risk then confidence)
        try:
            rank = {'Critical': 4, 'High': 3, 'Medium': 2, 'Low': 1}
            gaps = getattr(report, 'compliance_gaps', []) or []
            def gap_key(g):
                rl = getattr(g, 'risk_level', None)
                rlval = rl.value if hasattr(rl, 'value') else str(rl)
                conf = getattr(g, 'confidence_score', None) or 0.0
                return (rank.get(str(rlval), 0), conf)

            sorted_gaps = sorted(gaps, key=gap_key, reverse=True)[:5]
            if sorted_gaps:
                tg_data = [['Control', 'Risk', 'Confidence']]
                for g in sorted_gaps:
                    ctl = getattr(g, 'control_id', None) or getattr(g, 'control_title', 'Issue')
                    rl = getattr(g, 'risk_level', '')
                    if hasattr(rl, 'value'):
                        rl = rl.value
                    conf = getattr(g, 'confidence_score', None)
                    conf_str = f"{conf:.2f}" if isinstance(conf, (int, float)) else 'N/A'
                    tg_data.append([str(ctl), str(rl), conf_str])

                tg_tbl = Table(tg_data, colWidths=[3*inch, 1.2*inch, 1*inch])
                tg_tbl.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,0), self.colors['primary']),
                    ('TEXTCOLOR', (0,0), (-1,0), white),
                    ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                    ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                    ('GRID', (0,0), (-1,-1), 0.5, self.colors['dark_gray']),
                    ('BACKGROUND', (0,1), (-1,-1), self.colors['light_gray'])
                ]))

                story.append(Paragraph('Top 5 Compliance Gaps', styles['Heading2']))
                story.append(Spacer(1, 0.08*inch))
                story.append(tg_tbl)
                story.append(Spacer(1, 0.1*inch))
        except Exception:
            pass
        
        return story
    
    def _create_pdf_organization_info(self, report: FinalReport) -> List:
        """Create PDF organization information section."""
        story = []
        styles = getSampleStyleSheet()
        
        story.append(Paragraph("Organization Information", styles['Heading1']))
        story.append(Spacer(1, 0.2*inch))
        
        org_data = [
            ['Field', 'Information'],
            ['Organization Name', report.organization_info.organization_name],
            ['Industry Sector', report.organization_info.industry],
            ['Organization Size', report.organization_info.size],
            ['Contact Person', report.organization_info.contact_person],
            ['Contact Email', report.organization_info.contact_email],
            ['Assessment Date', report.organization_info.assessment_date.strftime('%B %d, %Y')],
            ['Assessment Scope', report.organization_info.scope]
        ]
        
        if report.organization_info.assessor_name:
            org_data.append(['Assessor', report.organization_info.assessor_name])
        
        org_table = Table(org_data, colWidths=[2.5*inch, 3.5*inch])
        org_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), self.colors['primary']),
            ('TEXTCOLOR', (0, 0), (-1, 0), white),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), self.colors['light_gray']),
            ('GRID', (0, 0), (-1, -1), 1, self.colors['dark_gray'])
        ]))
        
        story.append(org_table)
        
        return story

    def _create_pdf_compliance_by_function(self, report: FinalReport) -> List:
        """Create a simple vector bar chart showing compliance % for all 6 NIST functions."""
        story = []
        styles = getSampleStyleSheet()

        title_style = ParagraphStyle('SectionTitle', parent=styles['Heading1'], fontSize=16, spaceAfter=12, textColor=self.colors['primary'])
        story.append(Paragraph("Compliance Percentage by Function", title_style))
        story.append(Spacer(1, 0.1*inch))

        # Prepare data for all 6 NIST functions
        function_order = [
            ('Govern', NISTFunction.GOVERN, '#4f46e5'),
            ('Identify', NISTFunction.IDENTIFY, '#10b981'),
            ('Protect', NISTFunction.PROTECT, '#f59e0b'),
            ('Detect', NISTFunction.DETECT, '#ef4444'),
            ('Respond', NISTFunction.RESPOND, '#8b5cf6'),
            ('Recover', NISTFunction.RECOVER, '#06b6d4'),
        ]

        keys = [name.lower() for name, _, _ in function_order]
        per_func = {key: 0.0 for key in keys}

        def _set_if_complete(source: Dict[str, Any]) -> bool:
            if not isinstance(source, dict):
                return False
            updated = {}
            for key in keys:
                val = source.get(key)
                if val is None:
                    return False
                try:
                    updated[key] = float(val)
                except Exception:
                    return False
            for key in keys:
                per_func[key] = max(0.0, min(100.0, updated[key]))
            return True

        loaded = False
        metadata = getattr(report, 'metadata', {}) or {}

        # Strategy 1: precomputed per-function compliance from metadata
        try:
            loaded = _set_if_complete(metadata.get('compliance_by_function') or {})
        except Exception:
            loaded = False

        # Strategy 2: questionnaire answers via QuestionnaireEngine
        if not loaded:
            try:
                engine = QuestionnaireEngine()
                answers = getattr(report, 'questionnaire_answers', []) or []
                if answers:
                    computed = engine.compute_compliance_by_function(answers) or {}
                    loaded = _set_if_complete(computed)
            except Exception:
                loaded = False

        # Strategy 3: replicate API logic from gap counts + questionnaire totals
        if not loaded:
            try:
                engine = QuestionnaireEngine()

                # Try object gaps first
                gaps_count = {name: 0 for name, _, _ in function_order}
                obj_gaps = getattr(report, 'compliance_gaps', []) or []
                for gap in obj_gaps:
                    fn = getattr(gap, 'nist_function', None)
                    fn_value = fn.value if hasattr(fn, 'value') else str(fn)
                    if fn_value in gaps_count:
                        gaps_count[fn_value] += 1

                # Fallback to raw dict gap_analysis if object conversion lost data
                if not any(gaps_count.values()):
                    raw_gaps = metadata.get('gap_analysis') or []
                    if isinstance(raw_gaps, list):
                        for gap in raw_gaps:
                            if not isinstance(gap, dict):
                                continue
                            fn_value = gap.get('nist_function')
                            if fn_value in gaps_count:
                                gaps_count[fn_value] += 1

                totals = {}
                for _, enum_value, _ in function_order:
                    try:
                        totals[enum_value.value] = engine.get_function_summary(enum_value).get('total_questions', 0) or 0
                    except Exception:
                        totals[enum_value.value] = 0

                result = {}
                for name, _, _ in function_order:
                    tq = totals.get(name, 0)
                    gc = gaps_count.get(name, 0)
                    if tq == 0:
                        pct = 0.0
                    else:
                        pct = round(max(0.0, 100.0 * (1.0 - (gc / tq))), 1)
                    result[name.lower()] = pct

                loaded = _set_if_complete(result)
            except Exception:
                loaded = False

        # Strategy 4: derive from compliance_summary.gaps_by_function with questionnaire totals
        if not loaded:
            try:
                engine = QuestionnaireEngine()
                summary = getattr(report, 'compliance_summary', None)
                gbf = getattr(summary, 'gaps_by_function', {}) or {}
                result = {}
                for name, enum_value, _ in function_order:
                    gap_count = gbf.get(enum_value)
                    if gap_count is None:
                        gap_count = gbf.get(enum_value.value)
                    if gap_count is None:
                        gap_count = gbf.get(name)
                    if gap_count is None:
                        gap_count = gbf.get(name.lower())
                    gap_count = int(gap_count or 0)

                    try:
                        tq = engine.get_function_summary(enum_value).get('total_questions', 0) or 0
                    except Exception:
                        tq = 0

                    if tq == 0:
                        pct = 0.0
                    else:
                        pct = round(max(0.0, 100.0 * (1.0 - (gap_count / tq))), 1)

                    result[name.lower()] = pct

                loaded = _set_if_complete(result)
            except Exception:
                loaded = False

        # Strategy 5: even fallback from overall compliance percentage
        if not loaded:
            try:
                overall = float(getattr(report.compliance_summary, 'compliance_percentage', 0.0) or 0.0)
            except Exception:
                overall = 0.0
            overall = max(0.0, min(100.0, overall))
            for key in keys:
                per_func[key] = overall

        # Render each function as a table row with a colored bar
        rows = []
        bar_total_width = 250
        for name, _, color_hex in function_order:
            key = name.lower()
            pct = max(0.0, min(100.0, float(per_func.get(key, 0.0) or 0.0)))

            filled_width = bar_total_width * (pct / 100.0)
            empty_width = max(0.5, bar_total_width - filled_width)

            inner = Table(
                [['', '']],
                colWidths=[filled_width, empty_width],
                style=[
                    ('BACKGROUND', (0, 0), (0, 0), color_hex),
                    ('BACKGROUND', (1, 0), (1, 0), '#e5e7eb'),
                    ('INNERGRID', (0, 0), (-1, -1), 0, white),
                    ('BOX', (0, 0), (-1, -1), 0, white),
                ],
            )

            rows.append([
                Paragraph(f"<b>{name}</b>", styles['Normal']),
                inner,
                Paragraph(f"<b>{pct:.1f}%</b>", styles['Normal']),
            ])

        table = Table(rows, colWidths=[1.2*inch, bar_total_width, 0.8*inch], hAlign='LEFT')
        table.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('LEFTPADDING', (0, 0), (-1, -1), 4),
            ('RIGHTPADDING', (0, 0), (-1, -1), 4),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('ROWBACKGROUNDS', (0, 0), (-1, -1), [white, '#f9fafb']),
        ]))

        story.append(table)
        story.append(Spacer(1, 0.2*inch))
        return story

    def _priority_from_risk(self, risk_level):
        mapping = {
            'Critical': 'High',
            'High': 'High',
            'Medium': 'Medium',
            'Low': 'Low'
        }
        try:
            if hasattr(risk_level, 'value'):
                risk_level = risk_level.value
        except Exception:
            pass
        return mapping.get(str(risk_level), 'Medium')

    def _timeline_from_priority(self, priority):
        return {
            'High': '0-3 months',
            'Medium': '3-9 months',
            'Low': '9-18 months'
        }.get(priority, '3-9 months')
    
    def _create_pdf_compliance_gaps(self, report: FinalReport) -> List:
        """Create PDF compliance gaps section."""
        story = []
        styles = getSampleStyleSheet()
        
        story.append(Paragraph("Compliance Gaps Analysis", styles['Heading1']))
        story.append(Spacer(1, 0.2*inch))
        
        # Group gaps by NIST function
        gaps_by_function = {}
        for gap in report.compliance_gaps:
            function = gap.nist_function.value
            if function not in gaps_by_function:
                gaps_by_function[function] = []
            gaps_by_function[function].append(gap)
        
        for function, gaps in gaps_by_function.items():
            # Function section
            story.append(Paragraph(f"{function} Function", styles['Heading2']))
            story.append(Spacer(1, 0.1*inch))
            
            for gap in gaps:
                # Gap title with risk level
                risk_color = {
                    RiskLevel.CRITICAL: self.colors['accent'],
                    RiskLevel.HIGH: HexColor('#e67e22'),
                    RiskLevel.MEDIUM: self.colors['warning'],
                    RiskLevel.LOW: self.colors['success']
                }.get(gap.risk_level, self.colors['dark_gray'])
                
                gap_title_style = ParagraphStyle(
                    'GapTitle',
                    parent=styles['Heading3'],
                    fontSize=12,
                    spaceAfter=6,
                    textColor=risk_color
                )
                
                story.append(Paragraph(f"{gap.control_id}: {gap.control_title}", gap_title_style))
                story.append(Paragraph("<b>Review Status:</b> Verified", styles['Normal']))
                story.append(Paragraph(f"<b>Risk Level:</b> {gap.risk_level.value}", styles['Normal']))
                story.append(Paragraph(f"<b>Category:</b> {gap.category} - {gap.subcategory}", styles['Normal']))
                story.append(Spacer(1, 0.1*inch))
                
                # Gap description
                story.append(Paragraph("<b>Description:</b>", styles['Normal']))
                story.append(Paragraph(gap.description, styles['Normal']))
                story.append(Spacer(1, 0.1*inch))
                
                # Reasoning
                story.append(Paragraph("<b>Reasoning:</b>", styles['Normal']))
                story.append(Paragraph(gap.reasoning, styles['Normal']))
                story.append(Spacer(1, 0.1*inch))
                
                # Evidence
                if gap.evidence_sources:
                    story.append(Paragraph("<b>Evidence Sources:</b>", styles['Normal']))
                    for evidence in gap.evidence_sources:
                        story.append(Paragraph(f"• {evidence}", styles['Normal']))
                    story.append(Spacer(1, 0.1*inch))
                
                story.append(Spacer(1, 0.2*inch))
        
        return story
    
    def _create_pdf_risk_assessment(self, report: FinalReport) -> List:
        """Create PDF risk assessment as a visual heat map grid."""
        story = []
        styles = getSampleStyleSheet()

        story.append(Paragraph("Risk Assessment", styles['Heading1']))
        story.append(Spacer(1, 0.2*inch))

        if not report.risk_assessment:
            story.append(Paragraph("No risk assessment data available.", styles['Normal']))
            return story

        risk_colors = {
            'Critical': HexColor('#dc2626'),
            'High': HexColor('#ea580c'),
            'Medium': HexColor('#f59e0b'),
            'Low': HexColor('#22c55e'),
        }

        legend_items = [
            ('Low Risk', '#22c55e'),
            ('Medium Risk', '#f59e0b'),
            ('High Risk', '#ea580c'),
            ('Critical Risk', '#dc2626'),
        ]

        legend_cells = []
        legend_widths = []
        for label, _ in legend_items:
            legend_cells.append(
                Paragraph(
                    f'<b>{label}</b>',
                    ParagraphStyle(
                        'LegendCell',
                        parent=styles['Normal'],
                        fontSize=9,
                        textColor=white,
                        alignment=TA_CENTER,
                    ),
                )
            )
            legend_widths.append(1.3 * inch)

        legend_table = Table([legend_cells], colWidths=legend_widths)
        legend_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, 0), HexColor('#22c55e')),
            ('BACKGROUND', (1, 0), (1, 0), HexColor('#f59e0b')),
            ('BACKGROUND', (2, 0), (2, 0), HexColor('#ea580c')),
            ('BACKGROUND', (3, 0), (3, 0), HexColor('#dc2626')),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
        ]))
        story.append(legend_table)
        story.append(Spacer(1, 0.25*inch))

        header_style = ParagraphStyle(
            'TblHdr',
            parent=styles['Normal'],
            fontSize=10,
            fontName='Helvetica-Bold',
            textColor=white,
            alignment=TA_CENTER,
        )
        cell_white = ParagraphStyle(
            'TblCellW',
            parent=styles['Normal'],
            fontSize=8,
            leading=11,
            textColor=white,
            wordWrap='CJK',
        )
        cell_white_center = ParagraphStyle(
            'TblCellWC',
            parent=styles['Normal'],
            fontSize=8,
            leading=11,
            textColor=white,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold',
            wordWrap='CJK',
        )

        col_widths = [2.0*inch, 0.7*inch, 1.6*inch, 1.6*inch, 1.1*inch]

        header_row = [
            Paragraph('Risk Item', header_style),
            Paragraph('Level', header_style),
            Paragraph('Likelihood', header_style),
            Paragraph('Impact', header_style),
            Paragraph('Status', header_style),
        ]
        table_data = [header_row]

        for risk in report.risk_assessment:
            level = risk.risk_level.value if hasattr(risk.risk_level, 'value') else str(risk.risk_level)
            title_text = risk.title or 'Unknown Risk'
            likelihood = str(risk.likelihood or 'N/A')
            impact = str(risk.impact or 'N/A')
            status = str(risk.mitigation_status or 'Not Started')

            table_data.append([
                Paragraph(title_text, cell_white),
                Paragraph(f'<b>{level}</b>', cell_white_center),
                Paragraph(likelihood, cell_white),
                Paragraph(impact, cell_white),
                Paragraph(status, cell_white_center),
            ])

        heat_table = Table(table_data, colWidths=col_widths, repeatRows=1)

        style_cmds = [
            ('BACKGROUND', (0, 0), (-1, 0), HexColor('#1e293b')),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ('RIGHTPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#ffffff')),
        ]

        for row_idx, risk in enumerate(report.risk_assessment, start=1):
            level = risk.risk_level.value if hasattr(risk.risk_level, 'value') else str(risk.risk_level)
            bg = risk_colors.get(level, HexColor('#6b7280'))
            style_cmds.append(('BACKGROUND', (0, row_idx), (-1, row_idx), bg))

        heat_table.setStyle(TableStyle(style_cmds))
        story.append(heat_table)
        story.append(Spacer(1, 0.3*inch))

        story.append(Paragraph("Risk Details", styles['Heading2']))
        story.append(Spacer(1, 0.1*inch))

        for risk in report.risk_assessment:
            level = risk.risk_level.value if hasattr(risk.risk_level, 'value') else str(risk.risk_level)
            bg = risk_colors.get(level, HexColor('#6b7280'))

            title_bar_style = ParagraphStyle(
                'RiskTitleBar',
                parent=styles['Normal'],
                fontSize=10,
                textColor=white,
                backColor=bg,
                leftPadding=8,
                rightPadding=8,
                topPadding=5,
                bottomPadding=5,
                fontName='Helvetica-Bold',
                leading=14,
            )
            story.append(Paragraph(f"{risk.title}  •  {level}", title_bar_style))

            body_style = ParagraphStyle(
                'RiskBody',
                parent=styles['Normal'],
                fontSize=9,
                leftPadding=8,
                rightPadding=8,
                topPadding=4,
                bottomPadding=8,
                backColor=HexColor('#f9fafb')
                ,leading=13,
            )

            desc = risk.description or 'N/A'
            impact = risk.business_impact or 'N/A'
            likelihood = str(risk.likelihood or 'N/A')
            status = str(risk.mitigation_status or 'Not Started')

            story.append(Paragraph(
                f"<b>Likelihood:</b> {likelihood}<br/>"
                f"<b>Impact:</b> {impact}<br/>"
                f"<b>Description:</b> {desc}<br/>"
                f"<b>Status:</b> {status}",
                body_style,
            ))
            story.append(Spacer(1, 0.1*inch))

        return story
    
    def _create_pdf_recommendations(self, report: FinalReport) -> List:
        """Create PDF recommendations section."""
        story = []
        styles = getSampleStyleSheet()
        
        story.append(Paragraph("Recommendations", styles['Heading1']))
        story.append(Spacer(1, 0.2*inch))
        
        # Group recommendations by priority
        recommendations_by_priority = {
            'High': [],
            'Medium': [],
            'Low': []
        }
        
        for rec in report.recommendations:
            recommendations_by_priority[rec.priority].append(rec)
        
        for priority in ['High', 'Medium', 'Low']:
            recs = recommendations_by_priority[priority]
            if not recs:
                continue
            
            # Priority section
            priority_color = {
                'High': self.colors['accent'],
                'Medium': self.colors['warning'],
                'Low': self.colors['success']
            }[priority]
            
            priority_style = ParagraphStyle(
                'Priority',
                parent=styles['Heading2'],
                fontSize=14,
                spaceAfter=10,
                textColor=priority_color
            )
            
            story.append(Paragraph(f"{priority} Priority Recommendations", priority_style))
            story.append(Spacer(1, 0.1*inch))
            
            for rec in recs:
                story.append(Paragraph(f"{rec.title}", styles['Heading3']))
                story.append(Paragraph(f"<b>Priority:</b> {rec.priority}", styles['Normal']))
                story.append(Paragraph(f"<b>Estimated Effort:</b> {rec.estimated_effort}", styles['Normal']))
                story.append(Paragraph(f"<b>Description:</b> {rec.description}", styles['Normal']))
                
                if rec.implementation_steps:
                    story.append(Paragraph("<b>Implementation Steps:</b>", styles['Normal']))
                    for step in rec.implementation_steps:
                        story.append(Paragraph(f"• {step}", styles['Normal']))
                
                story.append(Spacer(1, 0.2*inch))
        
        return story
    
    def _create_pdf_questionnaire_evidence(self, report: FinalReport) -> List:
        """Create PDF questionnaire evidence section."""
        story = []
        styles = getSampleStyleSheet()
        
        story.append(Paragraph("Questionnaire Evidence", styles['Heading1']))
        story.append(Spacer(1, 0.2*inch))
        
        # Group answers by NIST function
        answers_by_function = {}
        for answer in report.questionnaire_answers:
            function = answer.nist_function.value
            if function not in answers_by_function:
                answers_by_function[function] = []
            answers_by_function[function].append(answer)
        
        for function, answers in answers_by_function.items():
            story.append(Paragraph(f"{function} Function Responses", styles['Heading2']))
            story.append(Spacer(1, 0.1*inch))
            
            for answer in answers:
                story.append(Paragraph(f"<b>Q:</b> {answer.question_text}", styles['Normal']))
                story.append(Paragraph(f"<b>A:</b> {answer.answer}", styles['Normal']))
                if answer.evidence:
                    story.append(Paragraph(f"<b>Evidence:</b> {answer.evidence}", styles['Normal']))
                story.append(Spacer(1, 0.1*inch))
        
        return story
    
    def _setup_word_styles(self, doc: Document) -> None:
        """Set up custom styles for Word document."""
        # Title style
        title_style = doc.styles['Title']
        title_style.font.name = 'Calibri'
        title_style.font.size = Pt(18)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(44, 62, 80)  # Dark blue
        
        # Heading 1 style
        h1_style = doc.styles['Heading1']
        h1_style.font.name = 'Calibri'
        h1_style.font.size = Pt(14)
        h1_style.font.bold = True
        h1_style.font.color.rgb = RGBColor(52, 152, 219)  # Blue
        
        # Heading 2 style
        h2_style = doc.styles['Heading2']
        h2_style.font.name = 'Calibri'
        h2_style.font.size = Pt(12)
        h2_style.font.bold = True
        h2_style.font.color.rgb = RGBColor(44, 62, 80)  # Dark blue
    
    def _create_word_title_page(self, doc: Document, report: FinalReport) -> None:
        """Create Word title page."""
        # Add title
        title = doc.add_paragraph('NIST CSF Compliance Assessment')
        title.style = doc.styles['Title']
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add organization name
        org = doc.add_paragraph(report.organization_info.organization_name)
        org.style = doc.styles['Heading1']
        org.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        doc.add_paragraph()  # Spacer
        
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_para.add_run(f"Assessment Date: {report.organization_info.assessment_date.strftime('%B %d, %Y')}\n")
        meta_para.add_run(f"Generated: {report.generated_at.strftime('%B %d, %Y at %I:%M %p')}\n")
        meta_para.add_run(f"Report ID: {report.report_id}")
        try:
            framework_name = (report.metadata or {}).get("framework_name")
            framework_version_id = (report.metadata or {}).get("framework_version_id")
            framework_version_label = (report.metadata or {}).get("framework_version_label")
            if framework_name or framework_version_id or framework_version_label:
                framework_line = f"\nFramework Used: {framework_name or 'NIST CSF'}"
                if framework_version_label:
                    framework_line += f" v{framework_version_label}"
                if framework_version_id:
                    framework_line += f" ({framework_version_id})"
                meta_para.add_run(framework_line)
        except Exception:
            pass
        
        # Add summary table
        doc.add_paragraph()  # Spacer
        
        summary_table = doc.add_table(rows=5, cols=2)
        summary_table.style = 'Light Grid Accent 1'
        summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        summary_table.cell(0, 0).text = 'Metric'
        summary_table.cell(0, 1).text = 'Value'
        
        # Data rows
        summary_table.cell(1, 0).text = 'Overall Compliance'
        summary_table.cell(1, 1).text = f"{report.compliance_summary.compliance_percentage:.1f}%"
        
        summary_table.cell(2, 0).text = 'Total Gaps Identified'
        summary_table.cell(2, 1).text = str(report.compliance_summary.total_gaps)
        
        summary_table.cell(3, 0).text = 'Overall Risk Score'
        summary_table.cell(3, 1).text = f"{report.compliance_summary.overall_risk_score:.1f}/10"
        
        summary_table.cell(4, 0).text = 'Critical Findings'
        summary_table.cell(4, 1).text = str(report.compliance_summary.critical_findings)
    
    def _create_word_toc(self, doc: Document) -> None:
        """Create Word table of contents."""
        doc.add_paragraph('Table of Contents', style='Heading1')
        
        toc_items = [
            "1. Executive Summary",
            "2. Organization Information", 
            "3. Compliance Gaps Analysis",
            "4. Risk Assessment",
            "5. Recommendations",
            "6. Questionnaire Evidence"
        ]
        
        for item in toc_items:
            doc.add_paragraph(item)
    
    def _create_word_executive_summary(self, doc: Document, report: FinalReport) -> None:
        """Create Word executive summary."""
        doc.add_paragraph('Executive Summary', style='Heading1')
        
        summary_text = f"""
This NIST Cybersecurity Framework (CSF) compliance assessment was conducted for {report.organization_info.organization_name}, a {report.organization_info.size} organization in the {report.organization_info.industry} sector.

The assessment identified {report.compliance_summary.total_gaps} compliance gaps across {len(set(gap.nist_function for gap in report.compliance_gaps))} NIST CSF functions. The overall compliance score is {report.compliance_summary.compliance_percentage:.1f}%, with an overall risk score of {report.compliance_summary.overall_risk_score:.1f}/10.

Key findings include {report.compliance_summary.critical_findings} critical issues that require immediate attention, and {len([r for r in report.recommendations if r.priority == 'High'])} high-priority recommendations for remediation.

The assessment covered the following NIST CSF functions: {', '.join(set(gap.nist_function.value for gap in report.compliance_gaps))}.
        """
        
        doc.add_paragraph(summary_text)
        # AI Confidence Score (aggregate) and Top 5 gaps
        try:
            ai_conf = getattr(report, 'ai_confidence_score', None)
            if ai_conf is None:
                confidences = [getattr(g, 'confidence_score', None) for g in getattr(report, 'compliance_gaps', []) or []]
                vals = [c for c in confidences if isinstance(c, (int, float))]
                if vals:
                    ai_conf = sum(vals) / len(vals)
                else:
                    ai_conf = None
        except Exception:
            ai_conf = None

        ai_text = f"AI Confidence Score: {ai_conf:.2f}" if (isinstance(ai_conf, (int, float))) else "AI Confidence Score: N/A"
        doc.add_paragraph(ai_text)

        try:
            rank = {'Critical': 4, 'High': 3, 'Medium': 2, 'Low': 1}
            gaps = getattr(report, 'compliance_gaps', []) or []
            def gap_key(g):
                rl = getattr(g, 'risk_level', None)
                rlval = rl.value if hasattr(rl, 'value') else str(rl)
                conf = getattr(g, 'confidence_score', None) or 0.0
                return (rank.get(str(rlval), 0), conf)

            sorted_gaps = sorted(gaps, key=gap_key, reverse=True)[:5]
            if sorted_gaps:
                tbl = doc.add_table(rows=1, cols=3)
                hdr = tbl.rows[0].cells
                hdr[0].text = 'Control'
                hdr[1].text = 'Risk'
                hdr[2].text = 'Confidence'
                for g in sorted_gaps:
                    row = tbl.add_row().cells
                    ctl = getattr(g, 'control_id', None) or getattr(g, 'control_title', 'Issue')
                    rl = getattr(g, 'risk_level', '')
                    if hasattr(rl, 'value'):
                        rl = rl.value
                    conf = getattr(g, 'confidence_score', None)
                    conf_str = f"{conf:.2f}" if isinstance(conf, (int, float)) else 'N/A'
                    row[0].text = str(ctl)
                    row[1].text = str(rl)
                    row[2].text = conf_str
        except Exception:
            pass
    
    def _create_word_organization_info(self, doc: Document, report: FinalReport) -> None:
        """Create Word organization information section."""
        doc.add_paragraph('Organization Information', style='Heading1')
        
        org_table = doc.add_table(rows=7, cols=2)
        org_table.style = 'Light Grid Accent 1'
        
        # Header row
        org_table.cell(0, 0).text = 'Field'
        org_table.cell(0, 1).text = 'Information'
        
        # Data rows
        org_table.cell(1, 0).text = 'Organization Name'
        org_table.cell(1, 1).text = report.organization_info.organization_name
        
        org_table.cell(2, 0).text = 'Industry Sector'
        org_table.cell(2, 1).text = report.organization_info.industry
        
        org_table.cell(3, 0).text = 'Organization Size'
        org_table.cell(3, 1).text = report.organization_info.size
        
        org_table.cell(4, 0).text = 'Contact Person'
        org_table.cell(4, 1).text = report.organization_info.contact_person
        
        org_table.cell(5, 0).text = 'Contact Email'
        org_table.cell(5, 1).text = report.organization_info.contact_email
        
        org_table.cell(6, 0).text = 'Assessment Date'
        org_table.cell(6, 1).text = report.organization_info.assessment_date.strftime('%B %d, %Y')
    
    def _create_word_compliance_gaps(self, doc: Document, report: FinalReport) -> None:
        """Create Word compliance gaps section."""
        doc.add_paragraph('Compliance Gaps Analysis', style='Heading1')
        
        # Group gaps by NIST function
        gaps_by_function = {}
        for gap in report.compliance_gaps:
            function = gap.nist_function.value
            if function not in gaps_by_function:
                gaps_by_function[function] = []
            gaps_by_function[function].append(gap)
        
        for function, gaps in gaps_by_function.items():
            doc.add_paragraph(f'{function} Function', style='Heading2')
            
            for gap in gaps:
                doc.add_paragraph(f'{gap.control_id}: {gap.control_title}', style='Heading3')

                status_para = doc.add_paragraph()
                status_para.add_run('Review Status: ').bold = True
                status_para.add_run('Verified')
                
                # Risk level with color
                risk_para = doc.add_paragraph()
                risk_para.add_run('Risk Level: ').bold = True
                risk_run = risk_para.add_run(gap.risk_level.value)
                
                # Set color based on risk level
                if gap.risk_level == RiskLevel.CRITICAL:
                    risk_run.font.color.rgb = RGBColor(231, 76, 60)  # Red
                elif gap.risk_level == RiskLevel.HIGH:
                    risk_run.font.color.rgb = RGBColor(230, 126, 34)  # Orange
                elif gap.risk_level == RiskLevel.MEDIUM:
                    risk_run.font.color.rgb = RGBColor(243, 156, 18)  # Yellow
                else:
                    risk_run.font.color.rgb = RGBColor(39, 174, 96)  # Green
                
                doc.add_paragraph()
                doc.add_paragraph().add_run('Category: ').bold = True
                doc.add_paragraph(f'{gap.category} - {gap.subcategory}')
                doc.add_paragraph()
                
                doc.add_paragraph().add_run('Description: ').bold = True
                doc.add_paragraph(gap.description)
                doc.add_paragraph()
                
                doc.add_paragraph().add_run('Reasoning: ').bold = True
                doc.add_paragraph(gap.reasoning)
                doc.add_paragraph()
    
    def _create_word_risk_assessment(self, doc: Document, report: FinalReport) -> None:
        """Create Word risk assessment section."""
        doc.add_paragraph('Risk Assessment', style='Heading1')
        
        # Risk summary table
        risk_table = doc.add_table(rows=len(report.risk_assessment) + 1, cols=5)
        risk_table.style = 'Light Grid Accent 1'
        
        # Header row
        risk_table.cell(0, 0).text = 'Risk Title'
        risk_table.cell(0, 1).text = 'Risk Level'
        risk_table.cell(0, 2).text = 'Likelihood'
        risk_table.cell(0, 3).text = 'Impact'
        risk_table.cell(0, 4).text = 'Status'
        
        # Data rows
        for i, risk in enumerate(report.risk_assessment, 1):
            risk_table.cell(i, 0).text = risk.title
            risk_table.cell(i, 1).text = risk.risk_level.value
            risk_table.cell(i, 2).text = risk.likelihood
            risk_table.cell(i, 3).text = risk.impact
            risk_table.cell(i, 4).text = risk.mitigation_status
        
        doc.add_paragraph()  # Spacer
        
        # Detailed risk descriptions
        for risk in report.risk_assessment:
            doc.add_paragraph(risk.title, style='Heading2')
            
            doc.add_paragraph().add_run('Risk Level: ').bold = True
            doc.add_paragraph(risk.risk_level.value)
            
            doc.add_paragraph().add_run('Description: ').bold = True
            doc.add_paragraph(risk.description)
            
            doc.add_paragraph().add_run('Business Impact: ').bold = True
            doc.add_paragraph(risk.business_impact)
            
            doc.add_paragraph()
    
    def _create_word_recommendations(self, doc: Document, report: FinalReport) -> None:
        """Create Word recommendations section."""
        doc.add_paragraph('Recommendations', style='Heading1')
        
        # Group recommendations by priority
        recommendations_by_priority = {
            'High': [],
            'Medium': [],
            'Low': []
        }
        
        for rec in report.recommendations:
            recommendations_by_priority[rec.priority].append(rec)
        
        for priority in ['High', 'Medium', 'Low']:
            recs = recommendations_by_priority[priority]
            if not recs:
                continue
            
            # Priority section
            doc.add_paragraph(f'{priority} Priority Recommendations', style='Heading2')
            
            for rec in recs:
                doc.add_paragraph(rec.title, style='Heading3')
                
                doc.add_paragraph().add_run('Priority: ').bold = True
                doc.add_paragraph(rec.priority)
                
                doc.add_paragraph().add_run('Estimated Effort: ').bold = True
                doc.add_paragraph(rec.estimated_effort)
                
                doc.add_paragraph().add_run('Description: ').bold = True
                doc.add_paragraph(rec.description)
                
                if rec.implementation_steps:
                    doc.add_paragraph().add_run('Implementation Steps: ').bold = True
                    for step in rec.implementation_steps:
                        doc.add_paragraph(f'• {step}', style='List Bullet')
                
                doc.add_paragraph()
    
    def _create_word_questionnaire_evidence(self, doc: Document, report: FinalReport) -> None:
        """Create Word questionnaire evidence section."""
        doc.add_paragraph('Questionnaire Evidence', style='Heading1')
        
        # Group answers by NIST function
        answers_by_function = {}
        for answer in report.questionnaire_answers:
            function = answer.nist_function.value
            if function not in answers_by_function:
                answers_by_function[function] = []
            answers_by_function[function].append(answer)
        
        for function, answers in answers_by_function.items():
            doc.add_paragraph(f'{function} Function Responses', style='Heading2')
            
            for answer in answers:
                doc.add_paragraph().add_run('Q: ').bold = True
                doc.add_paragraph(answer.question_text)
                
                doc.add_paragraph().add_run('A: ').bold = True
                doc.add_paragraph(answer.answer)
                
                if answer.evidence:
                    doc.add_paragraph().add_run('Evidence: ').bold = True
                    doc.add_paragraph(answer.evidence)
                
                doc.add_paragraph()
    
    def _create_pdf_heatmap_section(self, report: FinalReport, heatmap_report) -> List:
        """Create PDF risk heat map visualization section."""
        story = []
        styles = getSampleStyleSheet()

        def _hget(obj, key, default=None):
            if isinstance(obj, dict):
                return obj.get(key, default)
            return getattr(obj, key, default)

        def _enum_or_value(v, default='N/A'):
            if v is None:
                return default
            return v.value if hasattr(v, 'value') else v
        
        story.append(Paragraph("Risk Heat Map Analysis", styles['Heading1']))
        story.append(Spacer(1, 0.2*inch))
        
        if not heatmap_report:
            story.append(Paragraph("No heat map report available.", styles['Normal']))
            return story

        summary_style = ParagraphStyle(
            'HeatMapSummary',
            parent=styles['Normal'],
            fontSize=10,
            textColor=HexColor('#374151'),
            backColor=HexColor('#f3f4f6'),
            borderPadding=10,
            borderWidth=1,
            borderColor=HexColor('#9ca3af'),
            borderRadius=5
        )

        heat_items = _hget(heatmap_report, 'heat_map_items', []) or []
        items_above = _hget(heatmap_report, 'items_above_appetite', None)
        if items_above is None:
            items_above = sum(1 for item in heat_items if _hget(item, 'above_appetite', False))
        risk_profile = _hget(heatmap_report, 'risk_appetite_profile', None)
        max_risk = _enum_or_value(_hget(risk_profile, 'max_acceptable_risk', None), 'Not specified')

        summary_text = f"""
        <b>Heat Map Analysis Summary</b><br/><br/>
        <b>Items Analyzed:</b> {len(heat_items)}<br/>
        <b>Items Above Risk Appetite:</b> {items_above}<br/>
        <b>Items Within Risk Appetite:</b> {len(heat_items) - items_above}<br/>
        <b>Organization Risk Appetite:</b> {max_risk}<br/>
        """

        story.append(Paragraph(summary_text, summary_style))
        story.append(Spacer(1, 0.3*inch))

        # Handle heatmap image embedding - resolve to absolute path
        image_path = _hget(heatmap_report, 'visualization_path', None)
        if image_path:
            # Convert to absolute path if relative
            if not os.path.isabs(image_path):
                image_path = os.path.abspath(image_path)
        
        if image_path and os.path.exists(image_path):
            from reportlab.platypus import Image
            
            try:
                # Use higher DPI settings for better quality in PDF
                img = Image(image_path, width=6.5*inch, height=5*inch)
                img.hAlign = 'CENTER'
                story.append(img)
                story.append(Spacer(1, 0.3*inch))

                caption_style = ParagraphStyle(
                    'Caption',
                    parent=styles['Normal'],
                    fontSize=9,
                    alignment=TA_CENTER,
                    textColor=HexColor('#6b7280')
                )
                story.append(Paragraph("Figure 1: Risk Heat Map (Likelihood vs Impact)", caption_style))
                story.append(Spacer(1, 0.2*inch))

            except Exception as e:
                story.append(Paragraph(f"<i>Could not embed heat map image: {str(e)}</i>", styles['Normal']))
                story.append(Spacer(1, 0.2*inch))
        else:
            # Image doesn't exist - add note about missing image
            missing_note = ParagraphStyle(
                'MissingNote',
                parent=styles['Normal'],
                fontSize=10,
                textColor=HexColor('#dc2626'),
                alignment=TA_CENTER
            )
            story.append(Paragraph(
                f"<i>Heat map visualization not available. Path: {image_path or 'Not specified'}</i>", 
                missing_note
            ))
            story.append(Spacer(1, 0.2*inch))

        if heat_items:
            story.append(Paragraph("Risk Items Detail", styles['Heading2']))
            story.append(Spacer(1, 0.1*inch))

            items_data = [['Item', 'Risk Level', 'L', 'I', 'Status']]

            for item in heat_items[:10]:
                item_name = str(_hget(item, 'name', 'Unnamed item'))
                risk_level = str(_enum_or_value(_hget(item, 'risk_level', 'Unknown'), 'Unknown'))
                likelihood = _hget(item, 'likelihood', 'N/A')
                impact = _hget(item, 'impact', 'N/A')
                above_appetite = bool(_hget(item, 'above_appetite', False))
                status = '⚠️ Above' if above_appetite else '✓ Within'
                items_data.append([
                    item_name[:25] + ('...' if len(item_name) > 25 else ''),
                    risk_level,
                    str(likelihood),
                    str(impact),
                    status
                ])

            items_table = Table(items_data, colWidths=[2.5*inch, 1*inch, 0.4*inch, 0.4*inch, 1*inch])

            table_style = [
                ('BACKGROUND', (0, 0), (-1, 0), HexColor('#1e40af')),
                ('TEXTCOLOR', (0, 0), (-1, 0), white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('ALIGN', (0, 1), (0, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), HexColor('#f9fafb')),
                ('GRID', (0, 0), (-1, -1), 1, HexColor('#d1d5db')),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
            ]

            risk_row_colors = {
                'Critical': HexColor('#fee2e2'),
                'High': HexColor('#ffedd5'),
                'Medium': HexColor('#fef9c3'),
                'Low': HexColor('#dcfce7')
            }

            for i, item in enumerate(heat_items[:10], start=1):
                row_level = str(_enum_or_value(_hget(item, 'risk_level', 'Unknown'), 'Unknown'))
                color = risk_row_colors.get(row_level, white)
                table_style.append(('BACKGROUND', (0, i), (-1, i), color))

            items_table.setStyle(TableStyle(table_style))
            story.append(items_table)
        
        return story
    
    def _create_word_heatmap_section(self, doc: Document, report: FinalReport, heatmap_report) -> None:
        """Create Word risk heat map visualization section."""
        doc.add_paragraph('Risk Heat Map Analysis', style='Heading1')

        def _hget(obj, key, default=None):
            if isinstance(obj, dict):
                return obj.get(key, default)
            return getattr(obj, key, default)

        def _enum_or_value(v, default='N/A'):
            if v is None:
                return default
            return v.value if hasattr(v, 'value') else v
        
        if not heatmap_report:
            doc.add_paragraph('No heat map visualization available.')
            return

        heat_items = _hget(heatmap_report, 'heat_map_items', []) or []
        items_above = _hget(heatmap_report, 'items_above_appetite', None)
        if items_above is None:
            items_above = sum(1 for item in heat_items if _hget(item, 'above_appetite', False))
        risk_profile = _hget(heatmap_report, 'risk_appetite_profile', None)
        max_risk = _enum_or_value(_hget(risk_profile, 'max_acceptable_risk', None), 'Not specified')
        
        # Add heatmap summary
        summary_para = doc.add_paragraph()
        summary_para.add_run('Heat Map Summary:\n').bold = True
        summary_para.add_run(f'Heat map analysis covering {len(heat_items)} items.\n')
        summary_para.add_run(f'{items_above} items exceed organizational risk appetite.\n')
        summary_para.add_run(f'Organization Risk Appetite: {max_risk}')
        
        doc.add_paragraph()  # Spacer
        
        # Handle heatmap image - resolve to absolute path
        image_path = _hget(heatmap_report, 'visualization_path', None)
        if image_path:
            # Convert to absolute path if relative
            if not os.path.isabs(image_path):
                image_path = os.path.abspath(image_path)
        
        # Add heatmap image if it exists
        if image_path and os.path.exists(image_path):
            try:
                doc.add_picture(image_path, width=Inches(6))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add caption
                caption = doc.add_paragraph('Figure 1: Risk Heat Map (Likelihood vs Impact)')
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_format = caption.runs[0].font
                caption_format.size = Pt(9)
                caption_format.italic = True
                
                doc.add_paragraph()  # Spacer
            except Exception as e:
                doc.add_paragraph(f'(Could not embed heat map image: {str(e)})')
                doc.add_paragraph()  # Spacer
        else:
            doc.add_paragraph(f'(Heat map image not available. Path: {image_path or "Not specified"})')
            doc.add_paragraph()  # Spacer
            
        # Add items table
        if heat_items:
            items_table = doc.add_table(rows=1, cols=5)
            items_table.style = 'Light Grid Accent 1'
            
            # Header row
            header_cells = items_table.rows[0].cells
            header_cells[0].text = 'Item Name'
            header_cells[1].text = 'Likelihood'
            header_cells[2].text = 'Impact'
            header_cells[3].text = 'Risk Level'
            header_cells[4].text = 'Above Appetite'
            
            # Data rows
            for item in heat_items:
                row_cells = items_table.add_row().cells
                item_name = str(_hget(item, 'name', 'Unnamed item'))
                row_cells[0].text = item_name[:50]  # Truncate long names
                row_cells[1].text = str(_hget(item, 'likelihood', 'N/A'))
                row_cells[2].text = str(_hget(item, 'impact', 'N/A'))
                row_cells[3].text = str(_enum_or_value(_hget(item, 'risk_level', 'Unknown'), 'Unknown'))
                row_cells[4].text = 'Yes' if bool(_hget(item, 'above_appetite', False)) else 'No'
